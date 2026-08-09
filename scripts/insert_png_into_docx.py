from pathlib import Path
from docx import Document
from docx.shared import Inches


src = Path("D:/FinalProj/AI_Integrated_E_Voting_Project_Report.docx")
out = Path("D:/FinalProj/AI_Integrated_E_Voting_Project_Report_png_check.docx")
img = Path("D:/FinalProj/diagrams/png/fig_4_0_system_flowchart.png")

doc = Document(src)
doc.add_page_break()
doc.add_paragraph("PNG Size Verification Insert")
run = doc.add_paragraph().add_run()
run.add_picture(str(img), width=Inches(6.3))
doc.add_paragraph("Inserted: Fig 4.0 PNG at 6.3 inches width")
doc.save(out)

v = Document(out)
last = v.inline_shapes[-1]
print("output", out)
print("inline_shapes", len(v.inline_shapes))
print("width_inches", round(last.width.inches, 3))
print("height_inches", round(last.height.inches, 3))

